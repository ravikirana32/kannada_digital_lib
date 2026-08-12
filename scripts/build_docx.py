from pathlib import Path
import json
from docx import Document
ROOT=Path(__file__).resolve().parents[1]; DATA=ROOT/"data/sarvajna/tripadis"; OUT=ROOT/"books/sarvajna/generated/docx/sarvajna-samagra.docx"
doc=Document(); doc.add_heading("ಸರ್ವಜ್ಞ ಸಮಗ್ರ ತ್ರಿಪದಿಗಳು",0); doc.add_paragraph("ಮೂಲ ಕನ್ನಡ • ಭಾವಾರ್ಥ • ಜೀವನ ಸಂದೇಶ")
rows=[]
for p in DATA.glob("SAR-*.json"):
    if p.name.endswith(".template.json"): continue
    d=json.loads(p.read_text(encoding="utf-8"))
    if d["review_status"] in {"approved","published"}: rows.append(d)
for d in sorted(rows,key=lambda x:x["tripadi_number"]):
    doc.add_heading(f"ತ್ರಿಪದಿ {d['tripadi_number']}",1); doc.add_paragraph(d["original"])
    doc.add_heading("ಭಾವಾರ್ಥ",2); doc.add_paragraph(d["bhavartha"])
    doc.add_heading("ಜೀವನ ಸಂದೇಶ",2); doc.add_paragraph(d.get("life_message",""))
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print("DOCX generated:",OUT)
